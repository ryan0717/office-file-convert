import builtins
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import batch_convert_v2 as converter


class StructuredDocxTests(unittest.TestCase):
    def test_docx_preserves_structure_and_metadata(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            doc = Document()
            doc.sections[0].header.paragraphs[0].text = "Confidential Header"
            doc.sections[0].footer.paragraphs[0].text = "Page Footer"

            doc.add_paragraph("Hidden TOC Entry", style=doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH))
            doc.add_heading("Main Title", level=1)
            doc.add_heading("Sub Section", level=2)
            doc.add_paragraph("First numbered item", style="List Number")
            doc.add_paragraph("Second numbered item", style="List Number")

            paragraph = doc.add_paragraph()
            paragraph.add_run("Important").bold = True
            paragraph.add_run(" and ")
            paragraph.add_run("underlined").underline = True

            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Alpha"
            table.cell(1, 1).text = "42"
            doc.save(path)

            markdown = converter.docx_to_markdown(str(path))

        self.assertIn("## 文档元信息", markdown)
        self.assertIn("Confidential Header", markdown)
        self.assertIn("# Main Title", markdown)
        self.assertIn("## Sub Section", markdown)
        self.assertIn("1. First numbered item", markdown)
        self.assertIn("2. Second numbered item", markdown)
        self.assertIn("**Important**", markdown)
        self.assertIn("<u>underlined</u>", markdown)
        self.assertIn("| Name | Value |", markdown)
        self.assertIn("| Alpha | 42 |", markdown)
        self.assertNotIn("Hidden TOC Entry", markdown)


class StructuredPdfTests(unittest.TestCase):
    def _make_pdf(self, path):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF is not installed")

        doc = fitz.open()
        for page_no in range(2):
            page = doc.new_page(width=595, height=842)
            page.insert_text((72, 35), "Repeated Header", fontsize=9)
            page.insert_text((290, 810), str(page_no + 1), fontsize=9)
            if page_no == 0:
                page.insert_text((72, 90), "1 Overview", fontsize=18)
                page.insert_text((72, 125), "This paragraph should remain.", fontsize=11)
                top = 180
            else:
                page.insert_text((72, 90), "2 Details", fontsize=18)
                page.insert_text((72, 125), "Another paragraph should remain.", fontsize=11)
                top = 180

            xs = [72, 220, 420]
            ys = [top, top + 30, top + 60]
            for x in xs:
                page.draw_line((x, ys[0]), (x, ys[-1]), color=(0, 0, 0), width=0.5)
            for y in ys:
                page.draw_line((xs[0], y), (xs[-1], y), color=(0, 0, 0), width=0.5)
            page.insert_text((82, top + 20), "Metric", fontsize=10)
            page.insert_text((230, top + 20), "Rule", fontsize=10)
            page.insert_text((82, top + 50), f"Item {page_no + 1}", fontsize=10)
            page.insert_text((230, top + 50), "Keep this table text", fontsize=10)
        doc.save(path)

    def test_pdfplumber_pipeline_removes_repeated_headers_and_keeps_table(self):
        try:
            import pdfplumber  # noqa: F401
        except ImportError:
            self.skipTest("pdfplumber is not installed")

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.pdf"
            self._make_pdf(path)
            markdown, title = converter._extract_pdf_with_pdfplumber_structured(str(path))
            public_markdown, _public_title = converter.pdf_to_markdown(str(path))

        self.assertIn("<!-- page: 1 -->", markdown)
        self.assertIn("# 1 Overview", markdown)
        self.assertIn("# 2 Details", markdown)
        self.assertIn("This paragraph should remain.", markdown)
        self.assertIn("Keep this table text", markdown)
        self.assertIn("<!-- page: 1 -->", public_markdown)
        self.assertIn("Keep this table text", public_markdown)
        self.assertNotIn("Repeated Header", markdown)
        self.assertNotRegex(markdown, r"(?m)^1$")
        self.assertTrue(title.startswith("1 Overview"))

    def test_pdf_fallback_without_pdfplumber_does_not_crash(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.pdf"
            self._make_pdf(path)

            real_import = builtins.__import__

            def fake_import(name, *args, **kwargs):
                if name == "pdfplumber":
                    raise ImportError("simulated missing pdfplumber")
                return real_import(name, *args, **kwargs)

            with mock.patch("builtins.__import__", side_effect=fake_import):
                markdown, _title = converter.pdf_to_markdown(str(path))

        self.assertTrue(markdown.strip())


class RegressionTests(unittest.TestCase):
    def test_temporary_office_files_are_ignored(self):
        self.assertTrue(converter.is_temporary_office_file("~$draft.doc"))
        self.assertTrue(converter.is_temporary_office_file("~$report.docx"))
        self.assertFalse(converter.is_temporary_office_file("draft.doc"))


if __name__ == "__main__":
    unittest.main()
